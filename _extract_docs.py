import docx
import os

files = [
    r"C:\Users\User\OneDrive\Attachments\Desktop\Sign-to-Text_Setup_Guide.docx",
    r"C:\Users\User\OneDrive\Attachments\Desktop\doc1_setup.docx",
    r"C:\Users\User\OneDrive\Attachments\Desktop\doc2_godot.docx",
    r"C:\Users\User\OneDrive\Attachments\Desktop\doc3_training.docx",
    r"C:\Users\User\OneDrive\Attachments\Desktop\Weekly_Report_Updated.docx",
]

for f in files:
    print("=" * 80)
    print("FILE:", f)
    print("=" * 80)
    if not os.path.exists(f):
        print("[MISSING]")
        continue
    try:
        d = docx.Document(f)
        paras = [p.text for p in d.paragraphs]
        for p in paras:
            print(p)
        # tables
        for ti, t in enumerate(d.tables):
            print(f"\n--- TABLE {ti} ---")
            for row in t.rows:
                print(" | ".join(cell.text for cell in row.cells))
    except Exception as e:
        print("[ERROR]", e)
    print()
