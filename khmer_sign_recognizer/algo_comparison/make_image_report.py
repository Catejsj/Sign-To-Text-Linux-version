"""Build the Sign Language MNIST comparison report as .docx.

    python algo_comparison/make_image_report.py

Reads algo_comparison/results_image/results.json plus the PNG charts written by
run_image_comparison.py.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DIR = Path(__file__).resolve().parent / "results_image"
R = json.loads((DIR / "results.json").read_text(encoding="utf-8"))

PRETTY = {"lda": "LDA", "svm": "SVM", "logreg": "Logistic Regression",
          "knn": "k-NN", "gboost": "Gradient Boosting", "mlp": "MLP",
          "rf": "Random Forest", "gru": "GRU (recurrent)"}
RES = R["results"]
algos = sorted(RES, key=lambda a: RES[a]["acc"], reverse=True)
best = R["best"]


def _shade(cell, fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear"); el.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(el)


def para(doc, text="", size=11, bold=False, italic=False, align=None,
         color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10); run.font.name = "Calibri"
        c.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                                     else WD_ALIGN_PARAGRAPH.CENTER)
        _shade(c, "D9E2F3")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(10); run.font.name = "Calibri"
            cells[i].paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                                                else WD_ALIGN_PARAGRAPH.CENTER)
            if ri % 2:
                _shade(cells[i], "F2F2F2")
    for r_ in t.rows:
        for i, w in enumerate(widths):
            r_.cells[i].width = Inches(w)
    return t


def figure(doc, name, width, cap):
    p = DIR / name
    if not p.exists():
        para(doc, f"[chart missing: {name}]", italic=True); return
    doc.add_picture(str(p), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, cap, size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         color="595959", space_after=14)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
for m in ("top", "bottom", "left", "right"):
    setattr(sec, f"{m}_margin", Inches(1))
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

para(doc, "Sign Language Recognition from Images", size=22, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "Comparison of Machine Learning Algorithms", size=15,
     align=WD_ALIGN_PARAGRAPH.CENTER, color="404040", space_after=2)
para(doc, "Dataset: Sign Language MNIST (Kaggle)", size=10,
     align=WD_ALIGN_PARAGRAPH.CENTER, color="767171", space_after=16)

doc.add_heading("1. Objective", level=1)
para(doc, "Compare machine-learning algorithms for sign-language letter "
          "recognition from images, training every algorithm on the same public "
          "dataset under identical conditions so that differences in accuracy "
          "reflect the algorithm rather than the data or the split.")

doc.add_heading("2. Dataset", level=1)
table(doc, ["Property", "Value"], [
    ["Source", "Sign Language MNIST — Kaggle (datamunge/sign-language-mnist)"],
    ["Task", f"Static hand-sign letter classification, {R['n_classes']} classes"],
    ["Classes", ", ".join(R["classes"])],
    ["Image format", "28 x 28 pixels, grayscale (784 features)"],
    ["Training images", f"{R['n_train']:,}"],
    ["Test images", f"{R['n_test']:,}"],
    ["Split", "Official train/test split as published (not re-shuffled)"],
    ["Preprocessing", "Pixel values scaled to the range 0-1"],
    ["Metrics", "Accuracy and macro-F1 (macro-F1 weights every letter equally)"],
], [1.8, 4.7])
para(doc, "")
para(doc, "The letters J and Z are excluded from the dataset because they are "
          "produced with motion and cannot be represented by a single still "
          "image — a useful reminder that static-image methods cannot express "
          "the full language.", italic=True)
figure(doc, "samples.png", 5.6, "Figure 1 — Example images from the dataset.")

doc.add_page_break()

doc.add_heading("3. Results", level=1)
figure(doc, "comparison_chart.png", 6.3,
       "Figure 2 — Accuracy and macro-F1 by algorithm on the held-out test set.")
rows = [[PRETTY.get(a, a), f"{RES[a]['acc']:.2f}", f"{RES[a]['f1']:.2f}"]
        for a in algos]
table(doc, ["Algorithm", "Accuracy (%)", "Macro F1 (%)"], rows, [3.1, 1.9, 1.7])
para(doc, "")
para(doc, f"{PRETTY.get(best, best)} achieved the highest accuracy "
          f"({RES[best]['acc']:.2f}%). Accuracy and macro-F1 are close for every "
          f"algorithm, which indicates the classes are reasonably balanced and "
          f"no single letter dominates the result.")
para(doc, "Training times are deliberately not reported: they depend on the "
          "machine each algorithm was run on (a laptop, a desktop, or a cloud "
          "runtime with a different number of CPU cores) and so are not "
          "comparable between algorithms in this study.", italic=True)

doc.add_heading("4. Error analysis", level=1)
figure(doc, f"confusion_{best}.png", 5.2,
       f"Figure 3 — Confusion matrix for {PRETTY.get(best, best)}.")
para(doc, "Errors concentrate on letters whose handshapes differ only slightly "
          "at 28 x 28 resolution. Increasing image resolution, or using a "
          "convolutional model that exploits spatial structure rather than "
          "treating pixels as independent features, would be the natural next "
          "step for these cases.")

doc.add_heading("5. Discussion", level=1)
for t in [
    f"{PRETTY.get(best, best)} performed best on this dataset "
    f"({RES[best]['acc']:.2f}% accuracy).",
    "Accuracy and macro-F1 are close for every algorithm, which indicates the "
    "classes are reasonably balanced and no single letter dominates the result.",
    "All algorithms here treat the 784 pixels as independent features, so none "
    "of them uses the spatial arrangement of the image. This is the main "
    "limitation of the approach and explains why the remaining errors involve "
    "visually similar handshapes.",
    "This dataset contains only static letters photographed under controlled "
    "conditions. Real signing involves motion, continuous transitions between "
    "signs, and variation in lighting, background and camera angle, so accuracy "
    "measured here is an upper bound rather than a deployment estimate.",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(t); r.font.size = Pt(11); r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)

doc.add_heading("6. Reproducing these results", level=1)
p = doc.add_paragraph()
r = p.add_run("python algo_comparison/run_image_comparison.py --data archive.zip")
r.font.name = "Consolas"; r.font.size = Pt(10)
para(doc, "One script downloads nothing and changes nothing: it reads the "
          "Kaggle archive, trains every algorithm on the official split, and "
          "writes both the charts and the underlying numbers.")

out = DIR / "Image_Algorithm_Comparison_Report.docx"
doc.save(out)
print("wrote", out)
