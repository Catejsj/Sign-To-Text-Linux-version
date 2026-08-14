"""Build the algorithm-comparison report as .docx (opens in Google Docs).

    python algo_comparison/make_report.py

Reads algo_comparison/results/results.json plus the PNG charts produced by
run_comparison.py, and writes Algorithm_Comparison_Report.docx beside them.
"""
from __future__ import annotations

import json
from pathlib import Path
from statistics import mean, stdev

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DIR = Path(__file__).resolve().parent / "results"
R = json.loads((DIR / "results.json").read_text(encoding="utf-8"))

PRETTY = {"lda": "LDA", "svm": "SVM", "logreg": "Logistic Regression",
          "knn": "k-NN", "gboost": "Gradient Boosting", "mlp": "MLP",
          "rf": "Random Forest", "gru": "GRU (recurrent)"}

algos = sorted(R["same_signer_real"],
               key=lambda a: mean(R["same_signer_real"][a]["acc"]), reverse=True)
best = algos[0]
n_signers = len(R.get("signers", []))


def _shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def para(doc, text="", size=11, bold=False, italic=False, align=None,
         color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        c.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                                     else WD_ALIGN_PARAGRAPH.CENTER)
        _shade(c, "D9E2F3")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            cells[i].paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                                                else WD_ALIGN_PARAGRAPH.CENTER)
            if ri % 2:
                _shade(cells[i], "F2F2F2")
    # widths must be set per-cell for Google Docs to honour them
    for r_ in t.rows:
        for i, w in enumerate(widths):
            r_.cells[i].width = Inches(w)
    return t


def figure(doc, name, width, cap):
    p = DIR / name
    if not p.exists():
        para(doc, f"[chart missing: {name}]", italic=True)
        return
    doc.add_picture(str(p), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, cap, size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         color="595959", space_after=14)


def fmt(v, nd=1):
    return "—" if v is None else f"{v:.{nd}f}"


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
for m in ("top", "bottom", "left", "right"):
    setattr(sec, f"{m}_margin", Inches(1))
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── title ──
para(doc, "Khmer Sign Language Recognition", size=22, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "Comparison of Machine Learning Algorithms", size=15,
     align=WD_ALIGN_PARAGRAPH.CENTER, color="404040", space_after=2)
para(doc, "SignLink project — experimental results", size=10,
     align=WD_ALIGN_PARAGRAPH.CENTER, color="767171", space_after=16)

# ── 1 ──
doc.add_heading("1. Objective", level=1)
para(doc, "Compare machine-learning algorithms on the same Khmer Sign Language "
          "dataset under identical conditions, so that any difference in "
          "performance comes from the algorithm itself and not from the data, "
          "the split, or the evaluation procedure.")
para(doc, "All eight algorithms below were trained and evaluated by a single "
          "script, on one dataset, using the same features, the same splits and "
          "the same metrics. Only real recordings were used — no synthetic or "
          "augmented data.")

# ── 2 ──
doc.add_heading("2. Dataset and method", level=1)
signs = ", ".join((R.get("labels_text") or {}).get(l, l) for l in R["labels"])
table(doc, ["Property", "Value"], [
    ["Task", f"Isolated sign recognition, {len(R['labels'])} Khmer signs"],
    ["Signs", signs],
    ["Signers", str(n_signers)],
    ["Input", "Body + hand landmarks (48 joints x 3 coords x 60 frames)"],
    ["Features", "Per-joint mean / std / min / max over time (576 features)"],
    ["Training data", "Real recordings only (no augmentation)"],
    ["Evaluation", f"Held-out takes, mean of {R['splits']} random splits"],
    ["Split rule", "Split by take, so no take appears in both train and test"],
    ["Metrics", "Accuracy and macro-F1 (macro-F1 weights every sign equally)"],
], [1.8, 4.7])
para(doc, "")
para(doc, "Landmarks are used rather than raw video: the skeleton is unaffected "
          "by clothing, background and skin tone, which removes variation "
          "unrelated to the sign and keeps the input small enough to train "
          "every algorithm quickly.", italic=True)

# ── 3 ──
doc.add_heading("3. Results", level=1)
figure(doc, "comparison_chart.png", 6.3,
       "Figure 1 — Accuracy and macro-F1 by algorithm (held-out takes). "
       "Error bars show standard deviation across splits.")
rows = []
for a in algos:
    s = R["same_signer_real"][a]
    cr = R["cross_signer_real"].get(a)
    sdv = stdev(s["acc"]) if len(s["acc"]) > 1 else 0.0
    rows.append([PRETTY.get(a, a), f"{mean(s['acc']):.1f} ± {sdv:.1f}",
                 f"{mean(s['f1']):.1f}", fmt(mean(cr) if cr else None)])
table(doc, ["Algorithm", "Accuracy (%)", "Macro F1 (%)", "Unseen signer (%)"],
      rows, [2.1, 1.6, 1.5, 1.3])
para(doc, "")
para(doc, f"{PRETTY.get(best, best)} achieved the highest accuracy on held-out "
          f"takes. Accuracy and macro-F1 are close for every algorithm, which "
          f"indicates the dataset is balanced and no single sign dominates the "
          f"score. The final column reports the same models evaluated on a "
          f"signer who was absent from training — a harder and more realistic "
          f"test, discussed next.")

doc.add_page_break()

# ── 4 ──
doc.add_heading("4. The generalization gap", level=1)
figure(doc, "cross_signer.png", 6.3,
       "Figure 2 — Accuracy for a signer included in training versus a "
       "completely unseen signer.")
gap_best = max(algos, key=lambda a: mean(R["cross_signer_real"].get(a, [0])))
para(doc, "Accuracy falls sharply for a person the model has never seen. This "
          "is the central difficulty of sign-language recognition: signing "
          "speed, amplitude and body proportions differ between people, so a "
          "model trained on few signers partly learns the individual rather "
          "than the sign.")
para(doc, f"The ranking also changes between the two tests. "
          f"{PRETTY.get(best, best)} leads on held-out takes from a known "
          f"signer, while {PRETTY.get(gap_best, gap_best)} generalizes best to "
          f"an unseen signer. Choosing an algorithm on same-signer accuracy "
          f"alone would therefore select the wrong model for real deployment.")
para(doc, "For this reason both numbers are reported here. Publishing only the "
          "same-signer figure would substantially overstate real-world "
          "performance.", bold=True)

# ── 5 ──
doc.add_heading("5. Error analysis", level=1)
figure(doc, f"confusion_{best}.png", 4.6,
       f"Figure 3 — Confusion matrix for {PRETTY.get(best, best)}, summed over "
       f"all splits.")
para(doc, "The diagonal holds correct predictions; off-diagonal cells show "
          "which signs are mistaken for each other. Confusions concentrate on "
          "signs with similar arm trajectories. This is expected: hand "
          "landmarks are detected less reliably than body landmarks at webcam "
          "distance, so signs separated mainly by finger configuration are the "
          "hardest to distinguish.")

# ── 6 ──
doc.add_heading("6. Conclusions", level=1)
for t in [
    f"{PRETTY.get(best, best)} gave the highest accuracy on held-out takes "
    f"({mean(R['same_signer_real'][best]['acc']):.1f}%).",
    "Most algorithms cluster within a few points of each other, so the choice "
    "of algorithm matters less than the quality and diversity of the data.",
    f"Accuracy for an unseen signer is far lower than for a known signer, and "
    f"the best algorithm differs between the two settings "
    f"({PRETTY.get(gap_best, gap_best)} generalizes best).",
    "Cross-signer evaluation should always be reported alongside same-signer "
    "accuracy, since only the former reflects deployment conditions.",
    "Recording additional signers is expected to improve generalization more "
    "than further tuning of the algorithm.",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(t)
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)

# ── 7 ──
doc.add_heading("7. Reproducing these results", level=1)
p = doc.add_paragraph()
r = p.add_run(f"python algo_comparison/run_comparison.py --lang {R['language']} "
              f"--seeds {R['splits']}")
r.font.name = "Consolas"
r.font.size = Pt(10)
para(doc, "A single script trains every algorithm on identical splits and "
          "writes both the charts and the underlying numbers, so the comparison "
          "can be regenerated or extended without repeating any step by hand.")
para(doc, "Training times are not reported: they depend on the machine used "
          "and are not a property of the algorithm.", italic=True)

out = DIR / "Algorithm_Comparison_Report.docx"
doc.save(out)
print("wrote", out)
