"""Task A: compare algorithms on the khmer_var recordings, write the .docx.

    python algo_comparison/run_var_experiment.py

Uses ONLY data/sequences_v2/khmer_var — the 12-take-per-sign grid recording.
Nothing from the khmer corpus is involved.

Scoring is a take-aware random split averaged over several seeds: accuracy,
macro precision, macro recall and macro-F1. A take's variants never straddle
the split.

Writes charts plus Task_A_Report.docx into algo_comparison/results_khmer_var/.
"""
from __future__ import annotations

import json
import sys
import warnings
from pathlib import Path

import numpy as np

warnings.filterwarnings("ignore")
ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

import matplotlib                                                  # noqa: E402
matplotlib.use("Agg")
import matplotlib.pyplot as plt                                    # noqa: E402
from matplotlib import font_manager                                # noqa: E402
from sklearn.metrics import (f1_score, accuracy_score, confusion_matrix,
                             precision_score, recall_score)        # noqa: E402

from src.v2.algorithms import registry, wrap                       # noqa: E402
from src.v2.baseline_data import _featurize                        # noqa: E402
from src.v2.dataset import discover_samples                        # noqa: E402
from src.v2.schema import Source, View                             # noqa: E402

LANG, GRID, SEEDS = "khmer_var", 12, 8
OUT = ROOT / "algo_comparison" / f"results_{LANG}"

KHMER = ROOT / "fonts" / "NotoSansKhmer-Regular.ttf"
if KHMER.exists():
    font_manager.fontManager.addfont(str(KHMER))
    KH = font_manager.FontProperties(fname=str(KHMER), size=9)
else:
    KH = None

BLUE, GREY, RED = "#4472C4", "#A6A6A6", "#C00000"


def load():
    X, y, variants = [], [], []
    for p, m in discover_samples(ROOT / "data" / "sequences_v2", language=LANG):
        if m.source is not Source.REAL or m.view is not View.CLEAN:
            continue
        X.append(_featurize(np.load(p).astype(np.float32), "summary"))
        y.append(m.label); variants.append(m.variant)
    if not X:
        sys.exit(f"no real takes in data/sequences_v2/{LANG}/")
    labels = sorted(set(y)); l2i = {l: i for i, l in enumerate(labels)}
    return (np.stack(X), np.array([l2i[t] for t in y]),
            np.array(variants), labels)


X, y, variants, LABELS = load()
slot = variants % GRID
texts = {}
lp = ROOT / "data" / "sequences_v2" / LANG / "labels.json"
if lp.exists():
    texts = json.loads(lp.read_text(encoding="utf-8"))
LABEL_TEXT = [texts.get(l, l) for l in LABELS]
print(f"{len(X)} takes · {len(LABELS)} signs · {GRID} takes per sign\n")


def splits():
    """Take-aware random splits, one per seed."""
    takes = np.unique(variants)
    for seed in range(SEEDS):
        rs = np.random.default_rng(seed)
        held = set(rs.choice(takes, size=max(1, len(takes) // 4),
                             replace=False).tolist())
        yield np.array([v in held for v in variants])


def score(factory):
    acc, pre, rec, f1 = [], [], [], []
    for ev in splits():
        m = wrap(factory()); m.fit(X[~ev], y[~ev])
        p = m.predict(X[ev])
        acc.append(accuracy_score(y[ev], p) * 100)
        pre.append(precision_score(y[ev], p, average="macro",
                                   zero_division=0) * 100)
        rec.append(recall_score(y[ev], p, average="macro",
                                zero_division=0) * 100)
        f1.append(f1_score(y[ev], p, average="macro") * 100)
    return {"acc": float(np.mean(acc)), "pre": float(np.mean(pre)),
            "rec": float(np.mean(rec)), "f1": float(np.mean(f1)),
            "sd": float(np.std(f1))}


table_reg, _ = registry()
ALGOS = [a for a in table_reg if a != "nb"]
results = {}
for algo in ALGOS:
    label, factory, origin = table_reg[algo]
    results[algo] = {"name": label, **score(factory)}
    print(f"  {label:<22} F1 {results[algo]['f1']:5.1f} "
          f"± {results[algo]['sd']:.1f}")

order = sorted(results, key=lambda a: results[a]["f1"], reverse=True)
BEST = order[0]

# ---- best algorithm in detail, pooled over the same splits
_l, best_factory, _o = table_reg[BEST]
cm = np.zeros((len(LABELS), len(LABELS)), dtype=int)
ys, ps = [], []
for ev in splits():
    m = wrap(best_factory()); m.fit(X[~ev], y[~ev])
    p = m.predict(X[ev])
    cm += confusion_matrix(y[ev], p, labels=range(len(LABELS)))
    ys.append(y[ev]); ps.append(p)
ys, ps = np.concatenate(ys), np.concatenate(ps)
per_sign = f1_score(ys, ps, average=None, labels=range(len(LABELS))) * 100

# ---- did the recording conditions matter?
CONDITIONS = {
    "Lighting": ("full light", "dim light", slot < 6),
    "Distance": ("near", "far", (slot % 6) < 3),
    "Position": ("middle/left", "right", (slot % 3) < 2),
}
cond_rows = []
for name, (a_lbl, b_lbl, mask) in CONDITIONS.items():
    m = wrap(best_factory()); m.fit(X[mask], y[mask])
    held = f1_score(y[~mask], m.predict(X[~mask]), average="macro") * 100
    cond_rows.append((name, a_lbl, b_lbl, held))
baseline_f1 = results[BEST]["f1"]

# The lighting result turns out to depend heavily on the algorithm, so check
# every one rather than generalising from the single best-scoring model.
light_mask = CONDITIONS["Lighting"][2]
light_all = {}
for algo in ALGOS:
    _n, fac, _o = table_reg[algo]
    m = wrap(fac()); m.fit(X[light_mask], y[light_mask])
    light_all[algo] = f1_score(y[~light_mask], m.predict(X[~light_mask]),
                               average="macro") * 100
light_best = max(light_all, key=light_all.get)


# ─────────────────────────────────────────────────────────── charts

OUT.mkdir(parents=True, exist_ok=True)


def chart_comparison():
    fig, ax = plt.subplots(figsize=(8.6, 4.3))
    names = [results[a]["name"] for a in order]
    f1s = [results[a]["f1"] for a in order]
    sds = [results[a]["sd"] for a in order]
    colours = [BLUE if a == BEST else GREY for a in order]
    ax.bar(range(len(order)), f1s, yerr=sds, capsize=4, color=colours)
    for i, (v, s) in enumerate(zip(f1s, sds)):
        ax.text(i, v + s + 1.5, f"{v:.1f}", ha="center", fontsize=9)
    ax.set_xticks(range(len(order)))
    ax.set_xticklabels(names, rotation=25, ha="right", fontsize=9)
    ax.set_ylabel("macro-F1 (%)"); ax.set_ylim(0, 105)
    ax.set_title(f"Algorithm comparison — macro-F1 over {SEEDS} splits",
                 fontsize=12)
    ax.grid(axis="y", alpha=.3)
    fig.tight_layout(); fig.savefig(OUT / "var_comparison.png", dpi=150)
    plt.close(fig)


def chart_metrics():
    fig, ax = plt.subplots(figsize=(8.6, 4.0))
    xs = np.arange(len(order)); w = 0.27
    for k, (key, lbl, c) in enumerate([("pre", "Precision", "#8FAADC"),
                                       ("rec", "Recall", "#70AD47"),
                                       ("f1", "macro-F1", BLUE)]):
        ax.bar(xs + (k - 1) * w, [results[a][key] for a in order], w,
               label=lbl, color=c)
    ax.set_xticks(xs)
    ax.set_xticklabels([results[a]["name"] for a in order], rotation=25,
                       ha="right", fontsize=9)
    ax.set_ylabel("%"); ax.set_ylim(0, 105)
    ax.set_title("Precision, recall and macro-F1", fontsize=12)
    ax.legend(fontsize=9); ax.grid(axis="y", alpha=.3)
    fig.tight_layout(); fig.savefig(OUT / "var_metrics.png", dpi=150)
    plt.close(fig)


def chart_confusion():
    fig, ax = plt.subplots(figsize=(6.0, 5.2))
    norm = cm / np.maximum(cm.sum(axis=1, keepdims=True), 1)
    ax.imshow(norm, cmap="Blues", vmin=0, vmax=1)
    ax.set_xticks(range(len(LABELS))); ax.set_yticks(range(len(LABELS)))
    ax.set_xticklabels(LABEL_TEXT, rotation=45, ha="right",
                       fontproperties=KH if KH else None)
    ax.set_yticklabels(LABEL_TEXT, fontproperties=KH if KH else None)
    for i in range(len(LABELS)):
        for j in range(len(LABELS)):
            if cm[i, j]:
                ax.text(j, i, cm[i, j], ha="center", va="center", fontsize=9,
                        color="white" if norm[i, j] > .5 else "black")
    ax.set_xlabel("predicted"); ax.set_ylabel("actual")
    ax.set_title(f"{results[BEST]['name']} — confusion matrix", fontsize=11)
    fig.tight_layout(); fig.savefig(OUT / "var_confusion.png", dpi=150)
    plt.close(fig)


def chart_per_sign():
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    idx = np.argsort(per_sign)[::-1]
    ax.bar(range(len(LABELS)), per_sign[idx],
           color=[RED if v < 70 else BLUE for v in per_sign[idx]])
    ax.set_xticks(range(len(LABELS)))
    ax.set_xticklabels([LABEL_TEXT[i] for i in idx],
                       fontproperties=KH if KH else None)
    for i, v in enumerate(per_sign[idx]):
        ax.text(i, v + 1.5, f"{v:.0f}", ha="center", fontsize=9)
    ax.set_ylabel("macro-F1 (%)"); ax.set_ylim(0, 105)
    ax.set_title(f"{results[BEST]['name']} — score per sign", fontsize=11)
    ax.grid(axis="y", alpha=.3)
    fig.tight_layout(); fig.savefig(OUT / "var_per_sign.png", dpi=150)
    plt.close(fig)


chart_comparison(); chart_metrics(); chart_confusion(); chart_per_sign()
(OUT / "results.json").write_text(json.dumps(
    {"n_takes": int(len(X)), "labels": LABELS, "seeds": SEEDS,
     "best": BEST, "results": {a: results[a] for a in order},
     "conditions": [{"axis": n, "trained_on": a, "tested_on": b, "f1": s}
                    for n, a, b, s in cond_rows]},
    indent=2, ensure_ascii=False), encoding="utf-8")
print(f"\ncharts -> {OUT}")


# ─────────────────────────────────────────────────────────── report

from docx import Document                                          # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT                     # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH                      # noqa: E402
from docx.oxml import OxmlElement                                  # noqa: E402
from docx.oxml.ns import qn                                        # noqa: E402
from docx.shared import Inches, Pt, RGBColor                       # noqa: E402

CENTER = WD_ALIGN_PARAGRAPH.CENTER


def _shade(cell, fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear"); el.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(el)


def para(doc, text="", size=11, bold=False, italic=False, align=None,
         color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)


def table(doc, headers, rows, widths, highlight=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10); r.font.name = "Calibri"
        c.paragraphs[0].alignment = CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
        _shade(c, "D9E2F3")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(10); r.font.name = "Calibri"
            if highlight is not None and ri == highlight:
                r.bold = True
            cells[i].paragraphs[0].alignment = (CENTER if i
                                                else WD_ALIGN_PARAGRAPH.LEFT)
            if highlight is not None and ri == highlight:
                _shade(cells[i], "FFF2CC")
            elif ri % 2:
                _shade(cells[i], "F2F2F2")
    for r_ in t.rows:
        for i, w in enumerate(widths):
            r_.cells[i].width = Inches(w)
    return t


def figure(doc, name, width, cap):
    p = OUT / name
    if not p.exists():
        para(doc, f"[chart missing: {name}]", italic=True); return
    doc.add_picture(str(p), width=Inches(width))
    doc.paragraphs[-1].alignment = CENTER
    para(doc, cap, size=9, italic=True, align=CENTER, color="595959",
         space_after=14)


doc = Document()
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = s.bottom_margin = Inches(0.8)

h = doc.add_heading("Khmer Sign Language — Algorithm Comparison", level=0)
h.alignment = CENTER
para(doc, f"{len(ALGOS)} algorithms on {len(X)} recordings of "
          f"{len(LABELS)} signs", size=12, italic=True, align=CENTER,
     color="595959", space_after=16)

b = results[BEST]
second = results[order[1]]

# ---- 1
doc.add_heading("1. Data and method", level=1)
para(doc, "Signs are recognised from pose landmarks rather than images. Each "
          "recording is 60 frames of 48 tracked joints — six body points plus "
          "21 landmarks on each hand — which is summarised into 576 features "
          "(the mean, standard deviation, minimum and maximum of every joint "
          "coordinate over the clip).")
table(doc, ["", "Detail"],
      [["Signs", f"{len(LABELS)} — " + ", ".join(LABEL_TEXT)],
       ["Recordings", f"{len(X)} takes, {GRID} per sign"],
       ["Features", "576 summary values per take"],
       ["Split", f"75% train / 25% test, split by take, averaged over "
                 f"{SEEDS} random splits"],
       ["Algorithms", f"{len(ALGOS)}, each with the same feature scaling"]],
      [1.4, 5.1])
para(doc)
para(doc, "The split is by take, not by sample, so no take appears on both "
          "sides. Every algorithm is seeded, so the numbers reproduce exactly.",
     size=10, italic=True, color="595959")

# ---- 2
doc.add_heading("2. Results", level=1)
para(doc, "Macro-averaged metrics, so every sign counts equally regardless of "
          "how many recordings it has. Precision is how often a predicted sign "
          "is correct; recall is how often an actual sign is found; macro-F1 "
          "balances the two and is the headline number.")

rows = [[results[a]["name"], f"{results[a]['acc']:.1f}%",
         f"{results[a]['pre']:.1f}%", f"{results[a]['rec']:.1f}%",
         f"{results[a]['f1']:.1f}%", f"± {results[a]['sd']:.1f}"]
        for a in order]
table(doc, ["Algorithm", "Accuracy", "Precision", "Recall", "macro-F1", "s.d."],
      rows, [1.7, 0.95, 0.95, 0.9, 0.95, 0.75], highlight=0)
para(doc)
figure(doc, "var_comparison.png", 6.4,
       "Figure 1 — macro-F1 by algorithm. Bars show the standard deviation "
       f"across the {SEEDS} splits.")
figure(doc, "var_metrics.png", 6.4,
       "Figure 2 — precision, recall and macro-F1 side by side.")

spread = results[order[0]]["f1"] - results[order[-1]]["f1"]
# Anything within one standard deviation of the leader is not distinguishable
# on 84 takes; saying otherwise would read more into the ranking than it holds.
tied = [a for a in order
        if results[order[0]]["f1"] - results[a]["f1"] <= results[a]["sd"]]
if len(tied) > 1:
    names = ", ".join(results[a]["name"] for a in tied)
    para(doc, f"{b['name']} scores highest at {b['f1']:.1f}% macro-F1, but the "
              f"top group — {names} — sits within one standard deviation of "
              f"each other. On {len(X)} recordings that is a statistical tie: "
              f"any of them is a defensible choice, and the ordering between "
              f"them would likely change with more data.")
else:
    para(doc, f"{b['name']} scores highest at {b['f1']:.1f}% macro-F1, clear of "
              f"{second['name']} at {second['f1']:.1f}% by more than the "
              f"run-to-run variation.")
para(doc, f"The genuine separation is between that leading group and the rest: "
          f"the field spans {spread:.0f} points from best to worst, and "
          f"{results[order[-1]]['name']} at {results[order[-1]]['f1']:.1f}% is "
          f"far outside any margin of error.")
para(doc, "Precision and recall stay close together for the stronger "
          "algorithms, which means errors are spread across signs rather than "
          "one sign being systematically over-predicted.")

# ---- 3
doc.add_heading(f"3. {b['name']} in detail", level=1)
figure(doc, "var_per_sign.png", 6.0,
       "Figure 3 — score per sign. Red marks signs below 70%.")
weak_idx = [i for i in range(len(LABELS)) if per_sign[i] < 70]
if weak_idx:
    names = ", ".join(f"{LABEL_TEXT[i]} ({per_sign[i]:.0f}%)"
                      for i in sorted(weak_idx, key=lambda i: per_sign[i]))
    para(doc, f"Weakest signs: {names}. These would benefit most from extra "
              f"recordings.")
else:
    para(doc, "No sign falls below 70% — performance is even across the set.")
figure(doc, "var_confusion.png", 4.5,
       "Figure 4 — confusion matrix. Rows are the true sign, columns the "
       "prediction; the diagonal is correct.")

off = cm.copy(); np.fill_diagonal(off, 0)
if off.max() > 0:
    i, j = np.unravel_index(off.argmax(), off.shape)
    para(doc, f"The most frequent single mistake is {LABEL_TEXT[i]} predicted "
              f"as {LABEL_TEXT[j]} ({off[i, j]} times).")

# ---- 4
doc.add_heading("4. Did the recording conditions matter?", level=1)
para(doc, "The recordings were made on a deliberate grid: each sign was "
          "performed under two lighting levels, at two distances from the "
          "camera and in three standing positions. That lets us ask whether "
          "those choices affect recognition at all.")
para(doc, "To test it, the model was trained on one setting and tested on the "
          "other — for example trained only on brightly lit takes and tested "
          "only on dim ones. If a setting made no difference, the score would "
          "stay near the "
          f"{b['f1']:.0f}% baseline from Section 2.")

table(doc, ["Trained on", "Tested on", "macro-F1", "vs baseline"],
      [[a_lbl, b_lbl, f"{s:.1f}%", f"{s - baseline_f1:+.1f}"]
       for _n, a_lbl, b_lbl, s in cond_rows],
      [1.6, 1.5, 1.2, 1.3])
para(doc)

worst = min(cond_rows, key=lambda r: r[3])
best_c = max(cond_rows, key=lambda r: r[3])
para(doc, f"Yes, but unevenly. Moving nearer or further from the camera, and "
          f"standing to one side, cost relatively little — the model still "
          f"reaches {best_c[3]:.0f}% on a position it never trained on. "
          f"Lighting is the exception: training only on bright recordings and "
          f"testing on dim ones drops the score to {worst[3]:.0f}%, "
          f"{baseline_f1 - worst[3]:.0f} points below baseline.")
para(doc, "The practical reading is that the system tolerates people standing "
          "in slightly different places, but not a change in how the room is "
          "lit. Recordings intended for real use should therefore cover the "
          "lighting the system will actually meet, rather than assuming one "
          "session generalises.")
para(doc, f"How badly lighting hurts depends strongly on the algorithm, and "
          f"not in the order Section 2 would suggest. {b['name']} is the most "
          f"accurate model overall but among the most lighting-sensitive, "
          f"falling to {light_all[BEST]:.0f}%. The most resilient is "
          f"{results[light_best]['name']} at {light_all[light_best]:.0f}%, "
          f"despite ranking lower on the standard comparison:")
table(doc, ["Algorithm", "macro-F1, trained bright / tested dim",
            "Rank in Section 2"],
      [[results[a]["name"], f"{light_all[a]:.1f}%", f"#{order.index(a) + 1}"]
       for a in sorted(light_all, key=light_all.get, reverse=True)],
      [1.9, 2.6, 1.5])
para(doc)
para(doc, "The tree-based models hold up best. They classify by asking whether "
          "individual measurements fall above or below thresholds, which "
          "survives a shift affecting part of the input, whereas the linear "
          "models combine every feature into one weighted sum and are thrown "
          "off when those features move together. If the system has to work "
          "under lighting it has not seen, that difference matters more than "
          "the ranking in Section 2.")
para(doc, "One caveat worth stating: the dim takes were always recorded after "
          "the bright ones, so tiredness or drift in how the signs were "
          "performed cannot be fully separated from the lighting itself. "
          "Recording half the session in the reverse order would settle that.",
     size=10, italic=True, color="595959")

# ---- 5
doc.add_heading("5. Conclusions", level=1)
if len(tied) > 1:
    bullet(doc, f"{b['name']} scores highest at {b['f1']:.1f}% macro-F1 "
                f"({b['acc']:.1f}% accuracy), though "
                f"{len(tied)} algorithms are tied within run-to-run variation.")
else:
    bullet(doc, f"{b['name']} is the best of the {len(ALGOS)} algorithms "
                f"tested, at {b['f1']:.1f}% macro-F1 ({b['acc']:.1f}% "
                f"accuracy).")
bullet(doc, f"The spread between best and worst is {spread:.0f} points, so the "
            f"choice of algorithm materially affects results on this dataset.")
bullet(doc, "Lighting affects recognition noticeably; distance and standing "
            "position affect it far less.")
bullet(doc, f"Robustness to lighting does not follow overall accuracy: "
            f"{results[light_best]['name']} keeps "
            f"{light_all[light_best]:.0f}% when the lighting changes while "
            f"{b['name']}, the most accurate model overall, keeps only "
            f"{light_all[BEST]:.0f}%.")
if weak_idx:
    bullet(doc, f"{LABEL_TEXT[min(weak_idx, key=lambda i: per_sign[i])]} is the "
                f"hardest sign and is the clearest target for more recordings.")
bullet(doc, "All recordings come from one signer, so these figures describe "
            "recognition for a known person. Performance for someone new is a "
            "separate question and needs recordings from more of the team.")

doc.add_heading("6. Reproducing this", level=1)
para(doc, "python algo_comparison/run_var_experiment.py", size=10,
     color="1F4E79")
para(doc, f"Averages {SEEDS} take-aware random splits with fixed seeds; the "
          f"output is identical on any machine.", size=9, italic=True,
     color="595959")

path = OUT / "Task_A_Report.docx"
doc.save(path)
print(f"report -> {path}")
